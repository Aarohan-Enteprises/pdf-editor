from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import Response
import subprocess
import tempfile
import os
import platform
import shutil
import logging
import time
import glob as glob_module

# PyMuPDF + python-docx for PDF to DOCX conversion (optional)
try:
    import fitz  # PyMuPDF
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(
    title="PDF2.in API",
    description="API for PDF compression and other operations",
    version="1.0.0",
    docs_url=None,
    redoc_url=None,
    openapi_url=None,
)

# Security constants
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB
PDF_MAGIC_BYTES = b'%PDF'


def get_ghostscript_command():
    """Get the Ghostscript command for the current platform."""
    if platform.system() == 'Windows':
        # Try common Windows Ghostscript names (check PATH first)
        for cmd in ['gswin64c', 'gswin32c', 'gs']:
            found = shutil.which(cmd)
            if found:
                return found
        # Check common installation paths
        common_paths = [
            r'C:\Program Files\gs\gs10.06.0\bin\gswin64c.exe',
            r'C:\Program Files\gs\gs10.05.0\bin\gswin64c.exe',
            r'C:\Program Files\gs\gs10.04.0\bin\gswin64c.exe',
            r'C:\Program Files\gs\gs10.03.1\bin\gswin64c.exe',
            r'C:\Program Files\gs\gs10.02.1\bin\gswin64c.exe',
            r'C:\Program Files\gs\gs10.01.2\bin\gswin64c.exe',
            r'C:\Program Files\gs\gs10.00.0\bin\gswin64c.exe',
            r'C:\Program Files (x86)\gs\gs10.06.0\bin\gswin32c.exe',
            r'C:\Program Files (x86)\gs\gs10.04.0\bin\gswin32c.exe',
        ]
        for path in common_paths:
            if os.path.exists(path):
                return path
        return None
    else:
        # Unix-like systems
        return shutil.which('gs')


def get_libreoffice_command():
    """Get the LibreOffice command for the current platform."""
    if platform.system() == 'Windows':
        # Check PATH first
        for cmd in ['soffice', 'soffice.exe']:
            found = shutil.which(cmd)
            if found:
                return found
        # Check common installation paths using glob for version flexibility
        search_patterns = [
            r'C:\Program Files\LibreOffice*\program\soffice.exe',
            r'C:\Program Files (x86)\LibreOffice*\program\soffice.exe',
        ]
        for pattern in search_patterns:
            matches = glob_module.glob(pattern)
            if matches:
                # Return the first match (typically the most recent version)
                return matches[0]
        return None
    else:
        # Unix-like systems
        return shutil.which('soffice') or shutil.which('libreoffice')

# CORS for Next.js frontend
# In production, set ALLOWED_ORIGINS env var (comma-separated)
allowed_origins = os.environ.get(
    "ALLOWED_ORIGINS",
    "http://localhost:3000"
).split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_methods=["POST", "GET", "OPTIONS"],
    allow_headers=["*"],
    expose_headers=["X-Original-Size", "X-Compressed-Size"],
)


@app.on_event("startup")
async def startup_event():
    gs_cmd = get_ghostscript_command()
    if gs_cmd:
        logger.info(f"Ghostscript found at: {gs_cmd}")
    else:
        logger.warning("Ghostscript not found! Compression will not work.")

    lo_cmd = get_libreoffice_command()
    if lo_cmd:
        logger.info(f"LibreOffice found at: {lo_cmd}")
    else:
        logger.warning("LibreOffice not found! DOCX to PDF conversion will not work.")

    if PYMUPDF_AVAILABLE:
        logger.info("PyMuPDF + python-docx available for PDF to DOCX conversion")
    else:
        logger.warning("PyMuPDF not installed. Install with: pip install pymupdf python-docx")

# Ghostscript quality presets
QUALITY_SETTINGS = {
    "low": "/screen",      # 72 dpi, smallest file
    "medium": "/ebook",    # 150 dpi, balanced
    "high": "/printer",    # 300 dpi, best quality
}


@app.post("/api/compress")
async def compress_pdf(
    file: UploadFile = File(...),
    quality: str = Form("medium")
):
    if quality not in QUALITY_SETTINGS:
        raise HTTPException(status_code=400, detail="Invalid quality setting")

    if not file.filename or not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="File must be a PDF")

    start_time = time.time()

    # Check if Ghostscript is available
    gs_cmd = get_ghostscript_command()
    logger.info(f"Ghostscript command: {gs_cmd}")
    if not gs_cmd:
        raise HTTPException(
            status_code=500,
            detail="Ghostscript is not installed. Please install Ghostscript from https://ghostscript.com/releases/gsdnld.html"
        )

    # Read uploaded file
    content = await file.read()

    # Validate file size
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 50MB")

    # Validate PDF magic bytes
    if not content.startswith(PDF_MAGIC_BYTES):
        raise HTTPException(status_code=400, detail="Invalid PDF file")
    read_time = time.time()
    logger.info(f"TIMING: File read completed in {read_time - start_time:.3f}s, size: {len(content)} bytes")

    # Create temp files for input and output
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as input_file:
        input_file.write(content)
        input_path = input_file.name

    output_path = input_path.replace('.pdf', '_compressed.pdf')
    write_time = time.time()
    logger.info(f"TIMING: Temp file written in {write_time - read_time:.3f}s")

    try:
        # Run Ghostscript compression
        gs_command = [
            gs_cmd,
            '-sDEVICE=pdfwrite',
            '-dCompatibilityLevel=1.4',
            f'-dPDFSETTINGS={QUALITY_SETTINGS[quality]}',
            '-dNOPAUSE',
            '-dQUIET',
            '-dBATCH',
            f'-sOutputFile={output_path}',
            input_path
        ]

        logger.info(f"Running command: {' '.join(gs_command)}")

        result = subprocess.run(
            gs_command,
            capture_output=True,
            text=True,
            timeout=120  # 2 minute timeout
        )

        logger.info(f"Ghostscript return code: {result.returncode}")
        if result.stderr:
            logger.info(f"Ghostscript stderr: {result.stderr}")
        if result.stdout:
            logger.info(f"Ghostscript stdout: {result.stdout}")

        if result.returncode != 0:
            raise HTTPException(
                status_code=500,
                detail=f"Ghostscript compression failed: {result.stderr}"
            )

        gs_time = time.time()
        logger.info(f"TIMING: Ghostscript completed in {gs_time - write_time:.3f}s")

        # Read compressed file
        with open(output_path, 'rb') as f:
            compressed_content = f.read()

        total_time = time.time()
        logger.info(f"TIMING: Total processing time: {total_time - start_time:.3f}s")

        # Return compressed PDF
        return Response(
            content=compressed_content,
            media_type='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="compressed_{file.filename}"',
                'X-Original-Size': str(len(content)),
                'X-Compressed-Size': str(len(compressed_content)),
            }
        )

    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=500, detail="Compression timed out")

    finally:
        # Cleanup temp files
        if os.path.exists(input_path):
            os.remove(input_path)
        if os.path.exists(output_path):
            os.remove(output_path)


@app.post("/api/lock")
async def lock_pdf(
    file: UploadFile = File(...),
    password: str = Form(...)
):
    if not file.filename or not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="File must be a PDF")

    if not password or len(password) < 1:
        raise HTTPException(status_code=400, detail="Password is required")

    # Read uploaded file
    content = await file.read()

    # Validate file size
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 50MB")

    # Validate PDF magic bytes
    if not content.startswith(PDF_MAGIC_BYTES):
        raise HTTPException(status_code=400, detail="Invalid PDF file")

    # Create temp files
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as input_file:
        input_file.write(content)
        input_path = input_file.name

    output_path = input_path.replace('.pdf', '_locked.pdf')

    try:
        # Use Ghostscript to encrypt the PDF
        gs_cmd = get_ghostscript_command()
        if not gs_cmd:
            raise HTTPException(status_code=500, detail="Ghostscript is not installed")

        gs_command = [
            gs_cmd,
            '-sDEVICE=pdfwrite',
            '-dCompatibilityLevel=1.4',
            '-dNOPAUSE',
            '-dQUIET',
            '-dBATCH',
            f'-sOwnerPassword={password}',
            f'-sUserPassword={password}',
            '-dEncryptionR=3',
            '-dKeyLength=128',
            '-dPermissions=-3904',
            f'-sOutputFile={output_path}',
            input_path
        ]

        result = subprocess.run(
            gs_command,
            capture_output=True,
            text=True,
            timeout=120
        )

        if result.returncode != 0:
            raise HTTPException(status_code=500, detail=f"Failed to lock PDF: {result.stderr}")

        with open(output_path, 'rb') as f:
            locked_content = f.read()

        return Response(
            content=locked_content,
            media_type='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="locked_{file.filename}"',
            }
        )

    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=500, detail="Operation timed out")

    finally:
        if os.path.exists(input_path):
            os.remove(input_path)
        if os.path.exists(output_path):
            os.remove(output_path)


@app.post("/api/unlock")
async def unlock_pdf(
    file: UploadFile = File(...),
    password: str = Form(...)
):
    if not file.filename or not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="File must be a PDF")

    if not password:
        raise HTTPException(status_code=400, detail="Password is required")

    # Read uploaded file
    content = await file.read()

    # Validate file size
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 50MB")

    # Validate PDF magic bytes
    if not content.startswith(PDF_MAGIC_BYTES):
        raise HTTPException(status_code=400, detail="Invalid PDF file")

    # Create temp files
    with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as input_file:
        input_file.write(content)
        input_path = input_file.name

    output_path = input_path.replace('.pdf', '_unlocked.pdf')

    try:
        # Use Ghostscript to decrypt the PDF
        gs_cmd = get_ghostscript_command()
        if not gs_cmd:
            raise HTTPException(status_code=500, detail="Ghostscript is not installed")

        gs_command = [
            gs_cmd,
            '-sDEVICE=pdfwrite',
            '-dCompatibilityLevel=1.4',
            '-dNOPAUSE',
            '-dQUIET',
            '-dBATCH',
            f'-sPDFPassword={password}',
            f'-sOutputFile={output_path}',
            input_path
        ]

        result = subprocess.run(
            gs_command,
            capture_output=True,
            text=True,
            timeout=120
        )

        logger.info(f"Unlock return code: {result.returncode}")
        if result.stdout:
            logger.info(f"Unlock stdout: {result.stdout}")
        if result.stderr:
            logger.info(f"Unlock stderr: {result.stderr}")

        # Check for password errors - Ghostscript may return 0 even on password failure
        # Check both stdout and stderr as Ghostscript output varies
        all_output = ((result.stdout or '') + (result.stderr or '')).lower()
        password_error_keywords = ['password did not work', 'cannot decrypt', 'password', 'decrypt', 'encrypted', 'invalid', 'authentication']

        if any(keyword in all_output for keyword in password_error_keywords):
            raise HTTPException(status_code=400, detail="Incorrect password")

        if result.returncode != 0:
            raise HTTPException(status_code=500, detail=f"Failed to unlock PDF: {result.stderr or 'Unknown error'}")

        # Verify output file was created and has content
        if not os.path.exists(output_path) or os.path.getsize(output_path) == 0:
            raise HTTPException(status_code=400, detail="Incorrect password")

        with open(output_path, 'rb') as f:
            unlocked_content = f.read()

        return Response(
            content=unlocked_content,
            media_type='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="unlocked_{file.filename}"',
            }
        )

    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=500, detail="Operation timed out")

    finally:
        if os.path.exists(input_path):
            os.remove(input_path)
        if os.path.exists(output_path):
            os.remove(output_path)


@app.post("/api/docx-to-pdf")
async def convert_docx_to_pdf(file: UploadFile = File(...)):
    """Convert a DOCX file to PDF using LibreOffice."""
    # Validate file extension
    if not file.filename or not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="File must be a DOCX document")

    start_time = time.time()

    # Check if LibreOffice is available
    lo_cmd = get_libreoffice_command()
    logger.info(f"LibreOffice command: {lo_cmd}")
    if not lo_cmd:
        raise HTTPException(
            status_code=500,
            detail="LibreOffice is not installed. Please install LibreOffice from https://www.libreoffice.org/download/"
        )

    # Read uploaded file
    content = await file.read()

    # Validate file size
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 50MB")

    # Validate DOCX magic bytes (DOCX is a ZIP file starting with PK)
    if not content.startswith(b'PK'):
        raise HTTPException(status_code=400, detail="Invalid DOCX file")

    read_time = time.time()
    logger.info(f"TIMING: File read completed in {read_time - start_time:.3f}s, size: {len(content)} bytes")

    # Create temp directory for input and output
    temp_dir = tempfile.mkdtemp()
    input_path = os.path.join(temp_dir, 'input.docx')

    try:
        # Write input file
        with open(input_path, 'wb') as f:
            f.write(content)

        write_time = time.time()
        logger.info(f"TIMING: Temp file written in {write_time - read_time:.3f}s")

        # Run LibreOffice conversion
        # --headless: Run without UI
        # --convert-to pdf: Convert to PDF format
        # --outdir: Output directory
        lo_command = [
            lo_cmd,
            '--headless',
            '--convert-to', 'pdf',
            '--outdir', temp_dir,
            input_path
        ]

        logger.info(f"Running command: {' '.join(lo_command)}")

        result = subprocess.run(
            lo_command,
            capture_output=True,
            text=True,
            timeout=120  # 2 minute timeout
        )

        logger.info(f"LibreOffice return code: {result.returncode}")
        if result.stderr:
            logger.info(f"LibreOffice stderr: {result.stderr}")
        if result.stdout:
            logger.info(f"LibreOffice stdout: {result.stdout}")

        if result.returncode != 0:
            raise HTTPException(
                status_code=500,
                detail=f"LibreOffice conversion failed: {result.stderr or 'Unknown error'}"
            )

        lo_time = time.time()
        logger.info(f"TIMING: LibreOffice completed in {lo_time - write_time:.3f}s")

        # Find the output PDF file
        output_path = os.path.join(temp_dir, 'input.pdf')

        if not os.path.exists(output_path):
            raise HTTPException(
                status_code=500,
                detail="Conversion failed: output PDF not created"
            )

        # Read the converted PDF
        with open(output_path, 'rb') as f:
            pdf_content = f.read()

        total_time = time.time()
        logger.info(f"TIMING: Total processing time: {total_time - start_time:.3f}s")

        # Generate output filename from input
        output_filename = file.filename.rsplit('.', 1)[0] + '.pdf'

        return Response(
            content=pdf_content,
            media_type='application/pdf',
            headers={
                'Content-Disposition': f'attachment; filename="{output_filename}"',
            }
        )

    except subprocess.TimeoutExpired:
        raise HTTPException(status_code=500, detail="Conversion timed out")

    finally:
        # Cleanup temp directory and all files
        shutil.rmtree(temp_dir, ignore_errors=True)


def convert_pdf_to_docx_with_pymupdf(input_path: str, output_path: str) -> bool:
    """Convert PDF to DOCX using PyMuPDF + python-docx.

    Extracts text, images, and drawings from PDF and reconstructs in DOCX.
    Preserves formatting, layout, and horizontal/vertical lines.

    Returns True on success, False on failure.
    """
    if not PYMUPDF_AVAILABLE:
        return False

    try:
        # Open PDF
        pdf_doc = fitz.open(input_path)
        docx_doc = Document()

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]

            # Add page break between pages (except first)
            if page_num > 0:
                docx_doc.add_page_break()

            # Get page dimensions for scaling
            page_rect = page.rect
            page_width = page_rect.width
            page_height = page_rect.height

            # Extract text blocks with formatting
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

            for block in blocks:
                if block["type"] == 0:  # Text block
                    for line in block.get("lines", []):
                        para = docx_doc.add_paragraph()

                        for span in line.get("spans", []):
                            text = span.get("text", "")
                            if not text.strip():
                                continue

                            run = para.add_run(text)

                            # Apply font size
                            font_size = span.get("size", 11)
                            run.font.size = Pt(font_size)

                            # Apply font name
                            font_name = span.get("font", "")
                            if font_name:
                                run.font.name = font_name.split("+")[-1]  # Remove subset prefix

                            # Apply bold/italic based on font flags
                            flags = span.get("flags", 0)
                            run.font.bold = bool(flags & 2 ** 4)  # Bold flag
                            run.font.italic = bool(flags & 2 ** 1)  # Italic flag

                            # Apply text color
                            color = span.get("color", 0)
                            if color and color != 0:
                                r = (color >> 16) & 0xFF
                                g = (color >> 8) & 0xFF
                                b = color & 0xFF
                                run.font.color.rgb = RGBColor(r, g, b)

                elif block["type"] == 1:  # Image block
                    try:
                        # Extract image
                        bbox = block.get("bbox", [0, 0, 100, 100])
                        img_width = bbox[2] - bbox[0]
                        img_height = bbox[3] - bbox[1]

                        # Get image from block
                        xref = block.get("xref", 0)
                        if xref:
                            img_data = pdf_doc.extract_image(xref)
                            if img_data:
                                img_bytes = img_data["image"]
                                img_ext = img_data["ext"]

                                # Save image temporarily
                                img_temp_path = os.path.join(
                                    os.path.dirname(output_path),
                                    f"temp_img_{page_num}_{xref}.{img_ext}"
                                )
                                with open(img_temp_path, "wb") as img_file:
                                    img_file.write(img_bytes)

                                # Add image to document with appropriate size
                                max_width = Inches(6)  # Max width in document
                                scale = min(1.0, max_width / Pt(img_width).inches) if img_width > 0 else 1.0
                                docx_doc.add_picture(img_temp_path, width=Inches(img_width * scale / 72))

                                # Clean up temp image
                                os.remove(img_temp_path)
                    except Exception as img_error:
                        logger.warning(f"Failed to extract image: {img_error}")

            # Extract and add drawings (lines, rectangles, etc.)
            try:
                drawings = page.get_drawings()
                if drawings:
                    for drawing in drawings:
                        if drawing.get("items"):
                            for item in drawing["items"]:
                                if item[0] == "l":  # Line
                                    # Add horizontal rule for horizontal lines
                                    p1, p2 = item[1], item[2]
                                    if abs(p1.y - p2.y) < 2:  # Horizontal line
                                        para = docx_doc.add_paragraph()
                                        # Add bottom border to simulate horizontal line
                                        pBdr = OxmlElement('w:pBdr')
                                        bottom = OxmlElement('w:bottom')
                                        bottom.set(qn('w:val'), 'single')
                                        bottom.set(qn('w:sz'), '6')
                                        bottom.set(qn('w:space'), '1')
                                        bottom.set(qn('w:color'), '000000')
                                        pBdr.append(bottom)
                                        para._p.get_or_add_pPr().append(pBdr)
            except Exception as draw_error:
                logger.warning(f"Failed to extract drawings: {draw_error}")

        # Save the document
        docx_doc.save(output_path)
        pdf_doc.close()

        return os.path.exists(output_path)

    except Exception as e:
        logger.error(f"PyMuPDF conversion error: {str(e)}")
        return False


def fix_docx_margins(docx_path: str) -> bool:
    """Post-process DOCX to fix margins and alignment issues.

    LibreOffice converts PDFs using text frames (positioned boxes) rather than
    normal paragraphs. This function fixes both regular paragraphs and text frames.

    Returns True on success, False on failure.
    """
    if not PYMUPDF_AVAILABLE:
        return False

    try:
        from docx.shared import Inches, Pt, Twips
        from docx.oxml.ns import nsmap, qn
        import zipfile
        from lxml import etree

        doc = Document(docx_path)

        # Fix page margins (match typical PDF margins ~0.75 inch)
        for section in doc.sections:
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)

        # Fix paragraph indentation - remove all indents
        for para in doc.paragraphs:
            para.paragraph_format.left_indent = Twips(0)
            para.paragraph_format.right_indent = Twips(0)
            para.paragraph_format.first_line_indent = Twips(0)

        # Fix tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        para.paragraph_format.left_indent = Twips(0)
                        para.paragraph_format.right_indent = Twips(0)
                        para.paragraph_format.first_line_indent = Twips(0)

        doc.save(docx_path)

        # Now fix text frames via direct XML manipulation
        # LibreOffice creates text boxes with horizontal position offsets
        fix_docx_textframes(docx_path)

        logger.info("Fixed DOCX margins successfully")
        return True

    except Exception as e:
        logger.warning(f"Failed to fix margins: {e}")
        return False


def fix_docx_textframes(docx_path: str) -> bool:
    """Fix text frame positions in DOCX by directly manipulating XML.

    LibreOffice places text in positioned frames. This shifts all frames
    to start from the left margin.
    """
    import zipfile
    import shutil
    from lxml import etree

    try:
        # DOCX is a ZIP file - we need to modify document.xml
        temp_path = docx_path + '.tmp'

        # Word/DrawingML namespaces
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
            'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape',
            'wpg': 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup',
        }

        with zipfile.ZipFile(docx_path, 'r') as zip_in:
            with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zip_out:
                for item in zip_in.infolist():
                    data = zip_in.read(item.filename)

                    if item.filename == 'word/document.xml':
                        # Parse and modify the document XML
                        root = etree.fromstring(data)

                        # Find minimum X position across all frames to use as offset
                        min_x = None
                        positions = []

                        # Find all positionH elements (horizontal position of frames)
                        for pos_h in root.iter('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}positionH'):
                            pos_offset = pos_h.find('{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}posOffset')
                            if pos_offset is not None and pos_offset.text:
                                try:
                                    x_val = int(pos_offset.text)
                                    positions.append((pos_offset, x_val))
                                    if min_x is None or x_val < min_x:
                                        min_x = x_val
                                except ValueError:
                                    pass

                        # Shift all positions left by the minimum (so leftmost starts at 0)
                        if min_x and min_x > 0 and positions:
                            # Convert margin to EMUs (914400 EMUs = 1 inch, use 0.75 inch margin)
                            target_margin = int(0.75 * 914400)
                            shift = min_x - target_margin

                            if shift > 0:
                                for pos_offset, old_val in positions:
                                    new_val = old_val - shift
                                    pos_offset.text = str(max(0, new_val))
                                logger.info(f"Shifted {len(positions)} text frames left by {shift} EMUs")

                        data = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone='yes')

                    zip_out.writestr(item, data)

        # Replace original with modified
        shutil.move(temp_path, docx_path)
        return True

    except Exception as e:
        logger.warning(f"Failed to fix text frames: {e}")
        # Clean up temp file if it exists
        if os.path.exists(docx_path + '.tmp'):
            os.remove(docx_path + '.tmp')
        return False


def convert_pdf_to_docx_with_libreoffice(input_path: str, temp_dir: str) -> str | None:
    """Convert PDF to DOCX using LibreOffice.

    Tries draw_pdf_import first (better for visual elements),
    falls back to writer_pdf_import if needed.

    Returns the output path on success, None on failure.
    """
    lo_cmd = get_libreoffice_command()
    if not lo_cmd:
        return None

    # Use writer_pdf_import - draw_pdf_import produces worse results
    filters = ['writer_pdf_import']

    for infilter in filters:
        try:
            # Clean up any previous attempt
            output_path = os.path.join(temp_dir, 'input.docx')
            if os.path.exists(output_path):
                os.remove(output_path)

            lo_command = [
                lo_cmd,
                '--headless',
                f'--infilter={infilter}',
                '--convert-to', 'docx',
                '--outdir', temp_dir,
                input_path
            ]

            logger.info(f"Running command: {' '.join(lo_command)}")

            result = subprocess.run(
                lo_command,
                capture_output=True,
                text=True,
                timeout=120
            )

            logger.info(f"LibreOffice return code: {result.returncode}")
            if result.stderr:
                logger.info(f"LibreOffice stderr: {result.stderr}")
            if result.stdout:
                logger.info(f"LibreOffice stdout: {result.stdout}")

            if os.path.exists(output_path):
                logger.info(f"Conversion successful with filter: {infilter}")
                # Post-process to fix margins
                fix_docx_margins(output_path)
                return output_path

            logger.warning(f"Filter {infilter} did not produce output, trying next...")

        except subprocess.TimeoutExpired:
            logger.error(f"LibreOffice conversion timed out with filter: {infilter}")
            continue
        except Exception as e:
            logger.error(f"LibreOffice conversion error with filter {infilter}: {str(e)}")
            continue

    return None


def convert_pdf_to_docx_with_poppler(input_path: str, temp_dir: str) -> str | None:
    """Convert PDF to DOCX using Poppler (pdftohtml) + Pandoc pipeline.

    This two-step conversion often preserves layout better:
    1. pdftohtml converts PDF to HTML with positioning
    2. pandoc converts HTML to DOCX

    Returns the output path on success, None on failure.
    """
    try:
        # Check if pdftohtml and pandoc are available
        pdftohtml_check = subprocess.run(['which', 'pdftohtml'], capture_output=True)
        pandoc_check = subprocess.run(['which', 'pandoc'], capture_output=True)

        if pdftohtml_check.returncode != 0:
            logger.warning("pdftohtml not found")
            return None
        if pandoc_check.returncode != 0:
            logger.warning("pandoc not found")
            return None

        html_path = os.path.join(temp_dir, 'output.html')
        output_path = os.path.join(temp_dir, 'output.docx')

        # Step 1: Convert PDF to HTML using pdftohtml
        # -c: complex mode - preserves more layout/visual elements
        # -s: single HTML file (not multiple pages)
        # -noframes: don't generate frames
        pdftohtml_cmd = [
            'pdftohtml',
            '-c',           # complex mode for better layout
            '-s',           # single document
            '-noframes',    # no frames
            '-enc', 'UTF-8',
            input_path,
            html_path
        ]

        logger.info(f"Running pdftohtml: {' '.join(pdftohtml_cmd)}")
        result = subprocess.run(
            pdftohtml_cmd,
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode != 0:
            logger.error(f"pdftohtml failed: {result.stderr}")
            return None

        # pdftohtml creates filename.html (without the .html we passed)
        # or sometimes with -s it creates the exact filename
        actual_html = html_path
        if not os.path.exists(actual_html):
            # Try without extension
            base = os.path.join(temp_dir, 'output')
            for ext in ['.html', 's.html', '-html.html']:
                test_path = base + ext
                if os.path.exists(test_path):
                    actual_html = test_path
                    break

        if not os.path.exists(actual_html):
            # List files to debug
            files = os.listdir(temp_dir)
            logger.error(f"HTML file not found. Files in temp_dir: {files}")
            # Try to find any HTML file
            for f in files:
                if f.endswith('.html'):
                    actual_html = os.path.join(temp_dir, f)
                    break

        if not os.path.exists(actual_html):
            logger.error("No HTML output from pdftohtml")
            return None

        logger.info(f"HTML created at: {actual_html}")

        # Step 2: Convert HTML to DOCX using Pandoc
        pandoc_cmd = [
            'pandoc',
            actual_html,
            '-o', output_path,
            '--from', 'html',
            '--to', 'docx'
        ]

        logger.info(f"Running pandoc: {' '.join(pandoc_cmd)}")
        result = subprocess.run(
            pandoc_cmd,
            capture_output=True,
            text=True,
            timeout=60
        )

        if result.returncode != 0:
            logger.error(f"pandoc failed: {result.stderr}")
            return None

        if os.path.exists(output_path):
            logger.info("Poppler+Pandoc conversion successful")
            return output_path

        logger.error("DOCX output not created")
        return None

    except subprocess.TimeoutExpired:
        logger.error("Poppler/Pandoc conversion timed out")
        return None
    except Exception as e:
        logger.error(f"Poppler+Pandoc conversion error: {str(e)}")
        return None


def convert_pdf_to_docx_with_pdf2docx(input_path: str, output_path: str) -> bool:
    """Convert PDF to DOCX using pdf2docx library.

    pdf2docx is specifically designed for PDF to DOCX conversion and handles
    layout elements including lines, tables, and images.

    Returns True on success, False on failure.
    """
    try:
        from pdf2docx import Converter

        logger.info("Starting pdf2docx conversion...")
        cv = Converter(input_path)
        cv.convert(output_path)
        cv.close()

        if os.path.exists(output_path):
            logger.info("pdf2docx conversion successful")
            return True

        logger.error("pdf2docx did not create output file")
        return False

    except ImportError:
        logger.error("pdf2docx not installed")
        return False
    except Exception as e:
        logger.error(f"pdf2docx conversion error: {str(e)}")
        return False


# Check if pdf2docx is available
PDF2DOCX_AVAILABLE = False
try:
    from pdf2docx import Converter
    PDF2DOCX_AVAILABLE = True
    logger.info("pdf2docx is available")
except ImportError:
    logger.warning("pdf2docx not available")


@app.post("/api/pdf-to-docx")
async def convert_pdf_to_docx(
    file: UploadFile = File(...),
    engine: str = Form("pymupdf")
):
    """Convert a PDF file to DOCX.

    Args:
        file: The PDF file to convert
        engine: Conversion engine - "pymupdf" (default), "libreoffice", "poppler", or "pdf2docx"
    """
    # Validate engine parameter
    if engine not in ["pymupdf", "libreoffice", "poppler", "pdf2docx"]:
        raise HTTPException(status_code=400, detail="Invalid engine. Use 'pymupdf', 'libreoffice', 'poppler', or 'pdf2docx'")

    # Validate file extension
    if not file.filename or not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="File must be a PDF document")

    start_time = time.time()

    # Check engine availability
    use_pymupdf = engine == "pymupdf" and PYMUPDF_AVAILABLE
    use_libreoffice = engine == "libreoffice"
    use_poppler = engine == "poppler"
    use_pdf2docx = engine == "pdf2docx" and PDF2DOCX_AVAILABLE

    if use_pdf2docx and not PDF2DOCX_AVAILABLE:
        raise HTTPException(
            status_code=500,
            detail="pdf2docx not available on this server."
        )

    if use_libreoffice:
        lo_cmd = get_libreoffice_command()
        if not lo_cmd:
            raise HTTPException(
                status_code=500,
                detail="LibreOffice not available on this server."
            )

    # Read uploaded file
    content = await file.read()

    # Validate file size
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File too large. Maximum size is 50MB")

    # Validate PDF magic bytes
    if not content.startswith(PDF_MAGIC_BYTES):
        raise HTTPException(status_code=400, detail="Invalid PDF file")

    read_time = time.time()
    logger.info(f"TIMING: File read completed in {read_time - start_time:.3f}s, size: {len(content)} bytes")

    # Create temp directory for input and output
    temp_dir = tempfile.mkdtemp()
    input_path = os.path.join(temp_dir, 'input.pdf')
    output_path = os.path.join(temp_dir, 'output.docx')

    try:
        # Write input file
        with open(input_path, 'wb') as f:
            f.write(content)

        write_time = time.time()
        logger.info(f"TIMING: Temp file written in {write_time - read_time:.3f}s")

        docx_content = None
        engine_used = None

        # Try the requested engine
        if use_pymupdf:
            logger.info("Attempting conversion with PyMuPDF...")
            if convert_pdf_to_docx_with_pymupdf(input_path, output_path):
                engine_used = "pymupdf"
                with open(output_path, 'rb') as f:
                    docx_content = f.read()
                logger.info("PyMuPDF conversion successful")
            else:
                logger.warning("PyMuPDF conversion failed")

        elif use_poppler:
            logger.info("Attempting conversion with Poppler+Pandoc...")
            poppler_output = convert_pdf_to_docx_with_poppler(input_path, temp_dir)
            if poppler_output:
                engine_used = "poppler"
                with open(poppler_output, 'rb') as f:
                    docx_content = f.read()
                logger.info("Poppler+Pandoc conversion successful")
            else:
                logger.warning("Poppler+Pandoc conversion failed")

        elif use_libreoffice:
            logger.info("Attempting conversion with LibreOffice...")
            lo_output = convert_pdf_to_docx_with_libreoffice(input_path, temp_dir)
            if lo_output:
                engine_used = "libreoffice"
                with open(lo_output, 'rb') as f:
                    docx_content = f.read()
                logger.info("LibreOffice conversion successful")

        elif use_pdf2docx:
            logger.info("Attempting conversion with pdf2docx...")
            if convert_pdf_to_docx_with_pdf2docx(input_path, output_path):
                engine_used = "pdf2docx"
                with open(output_path, 'rb') as f:
                    docx_content = f.read()
                logger.info("pdf2docx conversion successful")
            else:
                logger.warning("pdf2docx conversion failed")

        if docx_content is None:
            raise HTTPException(
                status_code=500,
                detail="Conversion failed with all available engines"
            )

        total_time = time.time()
        logger.info(f"TIMING: Total processing time: {total_time - start_time:.3f}s (engine: {engine_used})")

        # Generate output filename from input
        output_filename = file.filename.rsplit('.', 1)[0] + '.docx'

        return Response(
            content=docx_content,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={
                'Content-Disposition': f'attachment; filename="{output_filename}"',
                'X-Conversion-Engine': engine_used,
            }
        )

    finally:
        # Cleanup temp directory and all files
        shutil.rmtree(temp_dir, ignore_errors=True)


@app.get("/api/engines")
async def get_available_engines():
    """Get available conversion engines and their status."""
    lo_cmd = get_libreoffice_command()
    gs_cmd = get_ghostscript_command()

    # Check if poppler and pandoc are available
    poppler_available = subprocess.run(['which', 'pdftohtml'], capture_output=True).returncode == 0
    pandoc_available = subprocess.run(['which', 'pandoc'], capture_output=True).returncode == 0

    return {
        "pdf_to_docx": {
            "pymupdf": {
                "available": PYMUPDF_AVAILABLE,
                "description": "Advanced conversion with text and images",
                "missing": [] if PYMUPDF_AVAILABLE else ["pymupdf", "python-docx"]
            },
            "libreoffice": {
                "available": lo_cmd is not None,
                "description": "Conversion using LibreOffice (preserves lines)",
                "path": lo_cmd
            },
            "poppler": {
                "available": poppler_available and pandoc_available,
                "description": "PDF→HTML→DOCX pipeline using Poppler and Pandoc (complex mode)",
                "missing": [] if (poppler_available and pandoc_available) else
                          (["poppler-utils"] if not poppler_available else []) +
                          (["pandoc"] if not pandoc_available else [])
            },
            "pdf2docx": {
                "available": PDF2DOCX_AVAILABLE,
                "description": "Dedicated PDF to DOCX converter (preserves layout, lines, tables)",
                "missing": [] if PDF2DOCX_AVAILABLE else ["pdf2docx"]
            }
        },
        "docx_to_pdf": {
            "libreoffice": {
                "available": lo_cmd is not None,
                "path": lo_cmd
            }
        },
        "compression": {
            "ghostscript": {
                "available": gs_cmd is not None,
                "path": gs_cmd
            }
        }
    }


@app.get("/health")
async def health_check():
    return {"status": "ok"}
